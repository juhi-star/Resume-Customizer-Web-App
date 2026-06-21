import io
import json
import os

from dotenv import load_dotenv

load_dotenv()

from typing import Optional
from fastapi import FastAPI, File, Form, HTTPException, UploadFile, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel
from docx import Document
from sqlalchemy.orm import Session
from starlette.types import ASGIApp, Scope, Receive, Send

from database import engine, get_db
from models import Base, User, ResumeHistory
from auth import (
    hash_password,
    verify_password,
    create_access_token,
    get_current_user,
)
from resume_parser import parse_resume, ResumeParseError
from llm_service import customize_resume, generate_cover_letter, LLMServiceError

Base.metadata.create_all(bind=engine)

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024
MIN_JD_LENGTH = 50
ALLOWED_EXTENSIONS = (".pdf", ".docx")

app = FastAPI(title="Resume Customizer API")


@app.options("/{path:path}")
async def options_route(path: str):
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "*",
            "Access-Control-Allow-Headers": "*",
            "Access-Control-Max-Age": "600",
        },
    )


class CORSASGIMiddleware:
    def __init__(self, app: ASGIApp):
        self.app = app

    async def __call__(self, scope: Scope, receive: Receive, send: Send):
        if scope["type"] != "http":
            await self.app(scope, receive, send)
            return

        async def send_wrapper(message):
            if message["type"] == "http.response.start":
                headers = dict(message.get("headers", []))
                headers[b"access-control-allow-origin"] = b"*"
                headers[b"access-control-allow-methods"] = b"*"
                headers[b"access-control-allow-headers"] = b"*"
                message["headers"] = [(k, v) for k, v in headers.items()]
            await send(message)

        await self.app(scope, receive, send_wrapper)


app.add_middleware(CORSASGIMiddleware)


@app.get("/api/health")
def health():
    return {"status": "ok"}


@app.post("/api/customize")
async def customize(
    job_description: str = Form(...),
    resume: UploadFile = File(...),
):
    jd = job_description.strip()
    if not jd:
        raise HTTPException(status_code=400, detail="Job description cannot be empty.")
    if len(jd) < MIN_JD_LENGTH:
        raise HTTPException(
            status_code=400,
            detail=f"Job description is too short. Please paste at least {MIN_JD_LENGTH} characters.",
        )

    filename = resume.filename or ""
    if not filename.lower().endswith(ALLOWED_EXTENSIONS):
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload a PDF or DOCX file.",
        )

    file_bytes = await resume.read()
    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="Uploaded resume file is empty.")
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=400,
            detail="Resume file is too large. Maximum allowed size is 5MB.",
        )

    try:
        resume_text = parse_resume(filename, file_bytes)
    except ResumeParseError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    try:
        result = customize_resume(jd, resume_text)
    except LLMServiceError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    return {
        "original_resume_text": resume_text,
        "customized": result,
    }


class ExperienceItem(BaseModel):
    title: str = ""
    company: str = ""
    duration: str = ""
    bullets: list[str] = []


class DownloadRequest(BaseModel):
    summary: str = ""
    relevant_skills: list[str] = []
    missing_keywords: list[str] = []
    experience: list[ExperienceItem] = []
    other_sections: str = ""
    candidate_name: str = "Candidate"


@app.post("/api/download")
def download_docx(payload: DownloadRequest):
    try:
        doc = Document()

        doc.add_heading(payload.candidate_name, level=0)

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(payload.summary or "")

        if payload.relevant_skills:
            doc.add_heading("Key Skills", level=1)
            doc.add_paragraph(", ".join(payload.relevant_skills))

        if payload.experience:
            doc.add_heading("Experience", level=1)
            for item in payload.experience:
                header_parts = [p for p in [item.title, item.company, item.duration] if p]
                p = doc.add_paragraph()
                run = p.add_run(" — ".join(header_parts))
                run.bold = True
                for bullet in item.bullets:
                    doc.add_paragraph(bullet, style="List Bullet")

        if payload.other_sections:
            doc.add_heading("Additional Information", level=1)
            for line in payload.other_sections.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        if payload.missing_keywords:
            doc.add_heading("Suggested Keywords to Consider Adding", level=1)
            doc.add_paragraph(", ".join(payload.missing_keywords))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=customized_resume.docx"},
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to generate document: {exc}") from exc


class CoverLetterRequest(BaseModel):
    job_description: str
    resume_text: str


@app.post("/api/cover-letter")
def cover_letter(payload: CoverLetterRequest):
    jd = payload.job_description.strip()
    if not jd or len(jd) < MIN_JD_LENGTH:
        raise HTTPException(status_code=400, detail="Job description is too short.")
    if not payload.resume_text.strip():
        raise HTTPException(status_code=400, detail="Resume text cannot be empty.")
    try:
        text = generate_cover_letter(jd, payload.resume_text)
        return {"cover_letter": text}
    except LLMServiceError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc



class RegisterRequest(BaseModel):
    email: str
    password: str


class LoginRequest(BaseModel):
    email: str
    password: str


def _user_response(user: User, token: str) -> dict:
    return {
        "access_token": token,
        "user": {"id": user.id, "email": user.email},
    }


@app.post("/api/auth/register")
def register(payload: RegisterRequest, db: Session = Depends(get_db)):
    if len(payload.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters.")
    existing = db.query(User).filter(User.email == payload.email).first()
    if existing:
        raise HTTPException(status_code=409, detail="An account with this email already exists.")
    user = User(email=payload.email, password_hash=hash_password(payload.password))
    db.add(user)
    db.commit()
    db.refresh(user)
    return _user_response(user, create_access_token(user.id))


@app.post("/api/auth/login")
def login(payload: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == payload.email).first()
    if not user or not verify_password(payload.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid email or password.")
    return _user_response(user, create_access_token(user.id))


@app.get("/api/auth/me")
def me(current_user: User = Depends(get_current_user)):
    return {"id": current_user.id, "email": current_user.email}



class SaveHistoryRequest(BaseModel):
    jd_snippet: str
    filename: str
    original_text: str
    customized_json: dict
    cover_letter: Optional[str] = None


@app.post("/api/history")
def save_history(
    payload: SaveHistoryRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = ResumeHistory(
        user_id=current_user.id,
        jd_snippet=payload.jd_snippet[:200],
        filename=payload.filename,
        original_text=payload.original_text,
        customized_json=json.dumps(payload.customized_json),
        cover_letter=payload.cover_letter,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"id": item.id, "created_at": item.created_at.isoformat()}


@app.get("/api/history")
def list_history(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    items = (
        db.query(ResumeHistory)
        .filter(ResumeHistory.user_id == current_user.id)
        .order_by(ResumeHistory.created_at.desc())
        .all()
    )
    return [
        {
            "id": item.id,
            "created_at": item.created_at.isoformat(),
            "jd_snippet": item.jd_snippet,
            "filename": item.filename,
        }
        for item in items
    ]


@app.get("/api/history/{item_id}")
def get_history_item(
    item_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = db.query(ResumeHistory).filter(
        ResumeHistory.id == item_id,
        ResumeHistory.user_id == current_user.id,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="History item not found.")
    return {
        "id": item.id,
        "created_at": item.created_at.isoformat(),
        "jd_snippet": item.jd_snippet,
        "filename": item.filename,
        "original_text": item.original_text,
        "customized": json.loads(item.customized_json),
        "cover_letter": item.cover_letter or "",
    }


@app.delete("/api/history/{item_id}")
def delete_history_item(
    item_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = db.query(ResumeHistory).filter(
        ResumeHistory.id == item_id,
        ResumeHistory.user_id == current_user.id,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="History item not found.")
    db.delete(item)
    db.commit()
    return {"ok": True}


class UpdateCoverLetterRequest(BaseModel):
    cover_letter: str


@app.patch("/api/history/{item_id}/cover-letter")
def update_cover_letter(
    item_id: int,
    payload: UpdateCoverLetterRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = db.query(ResumeHistory).filter(
        ResumeHistory.id == item_id,
        ResumeHistory.user_id == current_user.id,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="History item not found.")
    item.cover_letter = payload.cover_letter
    db.commit()
    return {"ok": True}
