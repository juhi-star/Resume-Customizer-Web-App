"""
resume_parser.py
Extracts plain text content from uploaded resume files (PDF or DOCX).
"""
import io
import pdfplumber
from docx import Document


class ResumeParseError(Exception):
    """Raised when a resume file cannot be parsed."""


def parse_pdf(file_bytes: bytes) -> str:
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n".join(text_parts).strip()
        if not text:
            raise ResumeParseError(
                "No extractable text found in PDF. The file may be a scanned image."
            )
        return text
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError(f"Failed to parse PDF: {exc}") from exc


def parse_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also pull text out of tables, since resumes sometimes use them
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        text = "\n".join(parts).strip()
        if not text:
            raise ResumeParseError("No extractable text found in DOCX file.")
        return text
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError(f"Failed to parse DOCX: {exc}") from exc


def parse_resume(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the correct parser based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif lower.endswith(".docx"):
        return parse_docx(file_bytes)
    else:
        raise ResumeParseError(
            "Unsupported file type. Please upload a PDF or DOCX file."
        )
